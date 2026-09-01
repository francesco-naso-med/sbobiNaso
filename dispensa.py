#!/usr/bin/env python3
"""
Generatore di dispense universitarie di medicina a partire da una sbobina grezza.

Pipeline:
  testo.txt -> chunking per paragrafi (~1 pagina) -> Gemini (rielaborazione, mai
  riassunto) -> Markdown -> Word nativo su base template.docx

Uso:
    export GEMINI_API_KEY="la-tua-chiave"
    python3 dispensa.py
    python3 dispensa.py --dry-run          # mostra solo come verrebbe diviso il testo
"""

from __future__ import annotations

import argparse
import json
import os
import random
import re
import shutil
import sys
import time
from dataclasses import dataclass
from pathlib import Path

from tqdm import tqdm

import sorgente
from sorgente import IMG_MARKER_RE

# ---------------------------------------------------------------------------
# CONFIGURAZIONE  (tutto quello che si tocca spesso sta qui)
# ---------------------------------------------------------------------------

INPUT_FILE = "testo.txt"
TEMPLATE_FILE = "template.docx"
OUTPUT_FILE = "dispensa.docx"
RAW_MARKDOWN_FILE = "dispensa_grezza.md"   # backup incrementale: salvato blocco per blocco

# Flash Lite: 500 richieste/giorno sul piano gratuito contro le 20 dei Flash pieni
# (tetto per progetto E per modello, verificato in AI Studio il 1/9/2026).
MODEL_NAME = "gemini-3.5-flash-lite"
# Rotazione: si passa al successivo quando un modello esaurisce la quota
# giornaliera, o se il suo nome non esiste (404). Il primo di riserva ha anche
# lui 500 richieste/giorno, gli altri 20.
FALLBACK_MODELS = ["gemini-3.1-flash-lite", "gemini-3.6-flash", "gemini-3.5-flash"]

CHUNK_TARGET_CHARS = 3500      # ~1 pagina A4 di testo pieno
CHUNK_MAX_CHARS = 5000         # limite duro: oltre si spezza comunque
CONTEXT_TAIL_CHARS = 700       # coda del blocco precedente passata come contesto

MAX_RETRIES = 5
BACKOFF_BASE = 2.0             # secondi: 2, 4, 8, 16, 32 (+ jitter)
BACKOFF_CAP = 60.0
TEMPERATURE = 0.3
SLEEP_BETWEEN_CALLS = 0.0      # pausa fra una chiamata e l'altra (quota al minuto)

FIRMA = "Dispensa rielaborata da sbobiNaso"   # piè di pagina di ogni pagina

# Rientri degli elenchi (cm). Il segno (pallino o numero) parte da
# LIST_INDENT_CM - LIST_HANGING_CM, il testo si allinea a LIST_INDENT_CM.
LIST_INDENT_CM = 1.25
LIST_HANGING_CM = 0.5

# Marcatore scritto nel backup Markdown: permette a --resume di sapere
# esattamente quanti blocchi sono gia' stati elaborati.
BLOCK_MARKER = "<!-- blocco {n} -->"
MARKER_RE = re.compile(r"^<!--\s*blocco\s+(\d+)\s*-->$")
COMMENT_RE = re.compile(r"^<!--.*-->$")

RETRY_DELAY_RE = re.compile(r"(?:retryDelay['\"]?\s*:\s*['\"]?|retry in )(\d+(?:\.\d+)?)s")


class QuotaExhaustedError(RuntimeError):
    """Quota GIORNALIERA esaurita: riprovare oggi e' inutile, non e' un errore transitorio."""


def _is_daily_quota(err: Exception) -> bool:
    msg = str(err).upper().replace(" ", "")
    return "429" in msg and ("PERDAY" in msg or "REQUESTSPERDAY" in msg)


def _server_retry_delay(err: Exception) -> float:
    """Rispetta il retryDelay suggerito dal server, quando c'e'."""
    match = RETRY_DELAY_RE.search(str(err))
    return float(match.group(1)) if match else 0.0

# ---------------------------------------------------------------------------
# PROMPT ACCADEMICO
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = """Sei un assistente accademico specializzato nella redazione di dispense \
universitarie di Medicina e Chirurgia.

Ricevi la trascrizione grezza (sbobina) di una lezione e la trasformi in testo da dispensa.

REGOLA ASSOLUTA — NON RIASSUMERE MAI.
Il tuo compito è RIELABORARE, non condensare. Ogni singolo contenuto informativo del testo \
originale deve sopravvivere nel testo finale: definizioni, numeri, percentuali, dosaggi, \
valori di laboratorio, nomi di farmaci, eponimi, classificazioni, criteri diagnostici, \
eccezioni, esempi clinici, casi citati dal docente, incisi e precisazioni. Se un dettaglio \
è nel testo di partenza, deve essere anche nel tuo output. In caso di dubbio, mantieni.
Il testo prodotto ha tipicamente lunghezza pari o superiore all'originale.

COSA DEVI FARE:
- Trasformare il parlato in italiano scritto accademico: periodi corretti, terminologia \
medica precisa, tono impersonale e professionale.
- Eliminare SOLO il rumore della trascrizione: intercalari ("ehm", "cioè", "no?"), \
ripetizioni identiche, false partenze, riferimenti d'aula puramente logistici \
("si vede in fondo?", "aprite la slide", "ci vediamo giovedì"), errori evidenti di \
trascrizione automatica (che vanno corretti, non rimossi, quando il senso è ricostruibile).
- Organizzare il contenuto in paragrafi coerenti, con titoli quando cambia argomento.
- Usare elenchi puntati o numerati quando il docente elenca voci, criteri, fasi o \
classificazioni.
- Evidenziare in grassetto i termini chiave: patologie, farmaci, strutture anatomiche, \
segni, valori soglia.

FORMATO DI OUTPUT (Markdown essenziale):
- `## Titolo` per le sezioni principali, `### Sottotitolo` per le sottosezioni.
- `- voce` per gli elenchi puntati, `1. voce` per quelli numerati.
- `**testo**` per il grassetto.
- Niente tabelle, niente blocchi di codice, niente `#` di primo livello.

IMMAGINI:
Nel testo puoi incontrare segnaposto come [[IMG:3]]: indicano una slide o una foto \
mostrata a lezione, che ti viene allegata insieme al testo. Quando ci sono:
- Riporta il segnaposto ESATTAMENTE com'è, da solo sulla sua riga, nel punto in cui \
l'immagine va mostrata. Non cambiarne il numero, non inventarne di nuovi, non ometterne.
- Usa quello che vedi per rendere esplicito il testo: al posto di rimandi vaghi come \
"in questa immagine si vede", scrivi che cosa si vede davvero ("la fotografia mostra \
lesioni simmetriche a entrambi gli arti inferiori").
- Se l'immagine contiene una tabella, uno schema o un elenco, trascrivine il contenuto \
nel testo: deve restare leggibile anche a chi stampa la dispensa in bianco e nero.

VINCOLI:
- Non aggiungere informazioni mediche non presenti nel testo di partenza.
- Non aggiungere introduzioni, conclusioni, commenti tuoi o meta-testo del tipo \
"in questo blocco", "in sintesi", "come abbiamo visto".
- Rispondi esclusivamente con il testo della dispensa, nient'altro."""

USER_TEMPLATE = """{context_block}### PORZIONE DI SBOBINA DA RIELABORARE (blocco {index} di {total})

{chunk}

### ISTRUZIONI
Rielabora integralmente il blocco qui sopra secondo le regole ricevute. Non riassumere: \
conserva tutti i dettagli. Non riprendere né ripetere il contesto precedente, serve solo \
a evitare ripetizioni e a mantenere la continuità del discorso."""

CONTEXT_TEMPLATE = """### CONTESTO — FINALE DEL BLOCCO PRECEDENTE (già scritto, NON riscriverlo)

{tail}

"""

CAPITOLI_PROMPT = """Ricevi l'elenco ordinato dei titoli delle sezioni di una dispensa \
universitaria. Raggruppa le sezioni CONSECUTIVE in capitoli tematicamente coerenti.

REGOLE:
- Rispetta l'ordine: un capitolo e' sempre un gruppo di sezioni consecutive.
- Punta a un totale compreso fra 6 e 15 capitoli per l'INTERA dispensa, qualunque \
sia il numero di sezioni: con molte sezioni i capitoli saranno semplicemente piu' \
corposi. Mai un capitolo per ogni sezione.
- Il titolo del capitolo deve essere breve (2-6 parole), in italiano, e descrivere \
l'argomento comune delle sezioni che contiene. Non ripetere pari pari il titolo di \
una sezione.
- La prima sezione appartiene sempre al primo capitolo.

Rispondi SOLO con un array JSON di oggetti con questi campi:
  "titolo": il titolo del capitolo
  "prima_sezione": il numero (1-based) della prima sezione del capitolo

Esempio: [{"titolo": "Fondamenti normativi", "prima_sezione": 1}, \
{"titolo": "Il consenso informato", "prima_sezione": 5}]"""


# ---------------------------------------------------------------------------
# 1. CHUNKING INTELLIGENTE
# ---------------------------------------------------------------------------


def split_sentences(text: str) -> list[str]:
    """Spezza un paragrafo troppo lungo sui confini di frase."""
    parts = re.split(r"(?<=[.!?;:])\s+", text)
    return [p for p in parts if p.strip()]


def chunk_text(text: str, target: int = CHUNK_TARGET_CHARS,
               hard_max: int = CHUNK_MAX_CHARS) -> list[str]:
    """Divide il testo in blocchi di ~1 pagina rispettando i confini di paragrafo.

    Un paragrafo non viene mai spezzato a metà, a meno che da solo non superi il
    limite duro: in quel caso si taglia sui confini di frase.
    """
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]

    units: list[str] = []
    for para in paragraphs:
        if len(para) <= hard_max:
            units.append(para)
            continue
        buffer = ""
        for sentence in split_sentences(para):
            if buffer and len(buffer) + len(sentence) + 1 > target:
                units.append(buffer.strip())
                buffer = sentence
            else:
                buffer = f"{buffer} {sentence}".strip()
        if buffer.strip():
            units.append(buffer.strip())

    chunks: list[str] = []
    current = ""
    for unit in units:
        candidate = f"{current}\n\n{unit}" if current else unit
        if current and len(candidate) > target:
            chunks.append(current)
            current = unit
        else:
            current = candidate
    if current.strip():
        chunks.append(current.strip())

    return chunks


# ---------------------------------------------------------------------------
# 2. CLIENT GEMINI + 3. RETRY CON BACKOFF ESPONENZIALE
# ---------------------------------------------------------------------------


def build_client(api_key: str | None = None):
    """Client Gemini. La chiave arriva esplicitamente (una per sessione) oppure
    dall'ambiente; mai dal codice.

    Passarla esplicitamente è indispensabile quando più persone usano la stessa
    istanza dell'app: scriverla in os.environ la renderebbe visibile a tutte le
    sessioni che condividono il processo.
    """
    try:
        from google import genai
    except ImportError:
        sys.exit("Manca la libreria ufficiale: pip install google-genai")

    api_key = api_key or os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        sys.exit(
            "Chiave API non trovata.\n"
            '  export GEMINI_API_KEY="la-tua-chiave"   (https://aistudio.google.com/apikey)'
        )
    return genai.Client(api_key=api_key)


def list_models() -> None:
    """Elenca gli ID API dei modelli utilizzabili: il nome mostrato in AI Studio
    ('Gemini 3.1 Flash Lite') non e' quello da scrivere nel codice."""
    # Il client va tenuto in una variabile: models.list() scarica le pagine man mano,
    # e un client temporaneo verrebbe chiuso a meta' iterazione.
    client = build_client()
    print("Modelli disponibili per la generazione di testo:\n")
    print(f"  {'ID API':<42} {'nome visualizzato':<34} contesto")
    try:
        models = list(client.models.list())
    except Exception as err:  # noqa: BLE001 - qui serve un messaggio leggibile, non un traceback
        if "API_KEY_INVALID" in str(err):
            sys.exit("Chiave API rifiutata da Google. Controlla GEMINI_API_KEY "
                     "(le chiavi di AI Studio iniziano con 'AIza').")
        sys.exit(f"Impossibile elencare i modelli: {err}")

    for model in models:
        actions = getattr(model, "supported_actions", None) or []
        if actions and "generateContent" not in actions:
            continue
        api_id = (model.name or "").replace("models/", "")
        limit = getattr(model, "input_token_limit", None)
        print(f"  {api_id:<42} {(model.display_name or '-'):<34} {limit or '-'}")


@dataclass
class GeminiWorker:
    model: str
    fallbacks: list[str]
    api_key: str | None = None      # se assente si usa quella dell'ambiente

    def __post_init__(self) -> None:
        from google.genai import types

        self._types = types
        self._client = build_client(self.api_key)
        self._pending_fallbacks = list(self.fallbacks)
        self._config = types.GenerateContentConfig(
            system_instruction=SYSTEM_PROMPT,
            temperature=TEMPERATURE,
            # Contenuto medico-accademico: senza questo, descrizioni di patologie,
            # farmaci e dosaggi possono far scattare i filtri di sicurezza.
            safety_settings=[
                types.SafetySetting(category=c, threshold="BLOCK_NONE")
                for c in (
                    "HARM_CATEGORY_HARASSMENT",
                    "HARM_CATEGORY_HATE_SPEECH",
                    "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                    "HARM_CATEGORY_DANGEROUS_CONTENT",
                )
            ],
        )

    # -- classificazione errori ------------------------------------------------

    @staticmethod
    def _is_transient(err: Exception) -> bool:
        msg = str(err).upper()
        markers = ("429", "500", "502", "503", "504", "RESOURCE_EXHAUSTED",
                   "UNAVAILABLE", "INTERNAL", "DEADLINE", "TIMEOUT", "OVERLOADED")
        return any(m in msg for m in markers)

    @staticmethod
    def _is_model_missing(err: Exception) -> bool:
        msg = str(err).upper()
        return ("404" in msg or "NOT_FOUND" in msg or "NOT FOUND" in msg) and "MODEL" in msg

    def _switch_model(self, reason: str = "non disponibile") -> bool:
        if not self._pending_fallbacks:
            return False
        old, self.model = self.model, self._pending_fallbacks.pop(0)
        tqdm.write(f"  ⚠️  '{old}' {reason} → passo a '{self.model}'")
        return True

    # -- chiamata con retry ----------------------------------------------------

    def raggruppa_capitoli(self, titoli: list[str]) -> dict[int, str]:
        """Una sola chiamata: restituisce {numero_sezione: titolo_capitolo}."""
        elenco = "\n".join(f"{i}. {t}" for i, t in enumerate(titoli, 1))
        config = self._types.GenerateContentConfig(
            system_instruction=CAPITOLI_PROMPT,
            temperature=0.2,
            response_mime_type="application/json",
            safety_settings=self._config.safety_settings,
        )
        response = self._client.models.generate_content(
            model=self.model, contents=elenco, config=config
        )
        gruppi = json.loads(response.text or "[]")

        capitoli: dict[int, str] = {}
        for gruppo in gruppi:
            inizio = int(gruppo["prima_sezione"])
            titolo = str(gruppo["titolo"]).strip()
            if titolo and 1 <= inizio <= len(titoli):
                capitoli[inizio] = titolo
        if capitoli:
            capitoli.setdefault(1, "Introduzione")   # il documento parte sempre da un capitolo
        return capitoli

    def _allegati(self, chunk: str, immagini: dict[int, Path]) -> list:
        """Le immagini citate dai segnaposto di questo blocco, come parti multimodali."""
        parti = []
        for numero in dict.fromkeys(int(n) for n in IMG_MARKER_RE.findall(chunk)):
            percorso = immagini.get(numero)
            if not percorso or not Path(percorso).is_file():
                continue
            suffisso = Path(percorso).suffix.lower().lstrip(".")
            mime = f"image/{'jpeg' if suffisso in ('jpg', 'jpeg') else suffisso}"
            parti.append(self._types.Part.from_bytes(
                data=Path(percorso).read_bytes(), mime_type=mime))
        return parti

    def elaborate(self, chunk: str, index: int, total: int, context_tail: str = "",
                  immagini: dict[int, Path] | None = None) -> str:
        context_block = CONTEXT_TEMPLATE.format(tail=context_tail) if context_tail else ""
        prompt = USER_TEMPLATE.format(
            context_block=context_block, index=index, total=total, chunk=chunk
        )
        contenuto = [prompt] + self._allegati(chunk, immagini or {})

        last_error: Exception | None = None
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                response = self._client.models.generate_content(
                    model=self.model, contents=contenuto, config=self._config
                )
                text = (response.text or "").strip()
                if text:
                    return text
                raise RuntimeError("risposta vuota dal modello (possibile blocco dei filtri)")

            except Exception as err:  # noqa: BLE001 - vogliamo classificare, non nascondere
                last_error = err

                if _is_daily_quota(err):
                    # Il tetto giornaliero e' per progetto E per modello: cambiare
                    # modello significa ripartire con un contatore intatto.
                    if self._switch_model("ha esaurito la quota giornaliera"):
                        continue
                    raise QuotaExhaustedError(
                        f"quota giornaliera esaurita su tutti i modelli disponibili "
                        f"({', '.join([self.model] + self.fallbacks)}) — fermo al blocco {index}"
                    ) from err

                if self._is_model_missing(err) and self._switch_model():
                    continue  # riprovo subito, senza consumare backoff

                if not self._is_transient(err) and not isinstance(err, RuntimeError):
                    raise RuntimeError(
                        f"Errore non recuperabile sul blocco {index}: {err}"
                    ) from err

                if attempt == MAX_RETRIES:
                    break

                own = BACKOFF_BASE ** attempt
                delay = min(max(own, _server_retry_delay(err)), BACKOFF_CAP) + random.uniform(0, 1)
                tqdm.write(
                    f"  ⏳ Blocco {index}: tentativo {attempt}/{MAX_RETRIES} fallito "
                    f"({type(err).__name__}). Riprovo tra {delay:.1f}s…"
                )
                time.sleep(delay)

        raise RuntimeError(
            f"Blocco {index}: falliti tutti i {MAX_RETRIES} tentativi. Ultimo errore: {last_error}"
        )


# ---------------------------------------------------------------------------
# 4. MARKDOWN -> WORD NATIVO (con template e fallback)
# ---------------------------------------------------------------------------

STYLE_CANDIDATES = {
    "h1": ["Titolo 1", "Heading 1"],
    "h2": ["Titolo 2", "Heading 2"],
    "h3": ["Titolo 3", "Heading 3"],
    "body": ["Corpo", "Corpo del testo", "Body Text", "Normale", "Normal"],
    "bullet": ["Elenco Puntato", "Elenco puntato", "List Bullet"],
    "number": ["Elenco Numerato", "Elenco numerato", "List Number"],
    # Stili di solo rientro: non portano il segno grafico, ma danno un elenco
    # decoroso quando il template non definisce quelli veri.
    "list_fallback": ["Paragrafo elenco", "List Paragraph"],
}

INLINE_RE = re.compile(r"(\*\*.+?\*\*|__.+?__|(?<!\*)\*(?!\s)[^*]+?\*|`[^`]+?`)")


def clear_body(doc) -> None:
    """Svuota il corpo di un documento mantenendone stili, margini e sezioni."""
    body = doc.element.body
    for child in list(body):
        if not child.tag.endswith("}sectPr"):   # la sectPr tiene margini/formato pagina
            body.remove(child)


def _add_field(paragraph, instruction: str, placeholder: str = "", dirty: bool = False) -> None:
    """Inserisce un campo Word (PAGE, TOC…) come sequenza di run.

    Word calcola il valore all'apertura: con dirty=True lo ricalcola da solo,
    altrimenti resta il testo segnaposto finché non si aggiorna il campo a mano.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    if dirty:
        begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for child in (begin, instr, separate, text, end):
        run = OxmlElement("w:r")
        run.append(child)
        paragraph._p.append(run)


class DocxBuilder:
    """Costruisce il .docx ereditando gli stili di template.docx, con fallback."""

    def __init__(self, template: str | None, cartella_immagini: Path | None = None):
        self.cartella_immagini = cartella_immagini
        from docx import Document

        self._Document = Document
        if template and Path(template).is_file():
            self.doc = Document(template)
            self._clear_template_body()
            self.template_used = True
        else:
            self.doc = Document()
            self.template_used = False

        available = {s.name for s in self.doc.styles}
        self.styles = {
            key: next((n for n in names if n in available), None)
            for key, names in STYLE_CANDIDATES.items()
        }

    def _clear_template_body(self) -> None:
        clear_body(self.doc)

    # -- helper ---------------------------------------------------------------

    @staticmethod
    def _disable_inherited_numbering(para) -> None:
        """Spegne la numerazione automatica ereditata dallo stile.

        Gli stili di elenco dei template Word portano un numId proprio: lasciarlo
        attivo significa un secondo pallino accanto al nostro, e — per gli elenchi
        numerati — un'unica numerazione continua per tutto il documento, che non
        riparte mai da 1. numId=0 la disattiva sul singolo paragrafo.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p_pr = para._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        for child in list(num_pr):
            num_pr.remove(child)
        for tag, value in (("w:ilvl", "0"), ("w:numId", "0")):
            el = OxmlElement(tag)
            el.set(qn("w:val"), value)
            num_pr.append(el)

    def _list_paragraph(self, kind: str, marker: str):
        """Paragrafo di elenco con il segno disegnato da noi.

        Il segno viene dal Markdown, quindi gli elenchi numerati ripartono da 1 a
        ogni elenco: e' la sorgente a dettare i numeri, non un contatore globale
        di Word. Dello stile del template restano font, corpo e spaziature.
        """
        from docx.shared import Cm

        para = self._paragraph(kind if self.styles.get(kind) else "list_fallback")
        self._disable_inherited_numbering(para)
        fmt = para.paragraph_format
        fmt.left_indent = Cm(LIST_INDENT_CM)
        fmt.first_line_indent = Cm(-LIST_HANGING_CM)
        para.add_run(marker)
        return para

    def _paragraph(self, style_key: str):
        style = self.styles.get(style_key)
        if style:
            try:
                return self.doc.add_paragraph(style=style)
            except KeyError:
                pass
        return self.doc.add_paragraph()   # fallback: stile di default del documento

    def _heading(self, text: str, level: int):
        key = {1: "h1", 2: "h2", 3: "h3"}.get(level, "h3")
        if self.styles.get(key):
            para = self._paragraph(key)
            self._add_runs(para, text)
            return
        # Fallback: nessuno stile titolo nel documento → grassetto dimensionato a mano
        from docx.shared import Pt

        para = self.doc.add_paragraph()
        run = para.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", text))
        run.bold = True
        run.font.size = Pt({1: 16, 2: 14, 3: 12}.get(level, 12))

    @staticmethod
    def _add_runs(paragraph, text: str) -> None:
        """Converte il Markdown inline (**grassetto**, *corsivo*, `code`) in run nativi."""
        for token in INLINE_RE.split(text):
            if not token:
                continue
            if (token.startswith("**") and token.endswith("**")) or (
                token.startswith("__") and token.endswith("__")
            ):
                paragraph.add_run(token[2:-2]).bold = True
            elif token.startswith("*") and token.endswith("*") and len(token) > 2:
                paragraph.add_run(token[1:-1]).italic = True
            elif token.startswith("`") and token.endswith("`") and len(token) > 2:
                paragraph.add_run(token[1:-1]).italic = True
            else:
                paragraph.add_run(token)

    # -- parser di blocco ------------------------------------------------------

    def add_image(self, numero: int) -> bool:
        """Inserisce l'immagine numero n, larga quanto la colonna di testo."""
        from docx.shared import Emu

        if not self.cartella_immagini:
            return False
        percorso = sorgente.percorso_immagine(self.cartella_immagini, numero)
        if not percorso:
            return False

        sezione = self.doc.sections[0]
        larghezza = Emu(sezione.page_width - sezione.left_margin - sezione.right_margin)
        try:
            self.doc.add_picture(str(percorso), width=larghezza)
        except Exception:  # noqa: BLE001 - un'immagine illeggibile non deve fermare la dispensa
            return False
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True

    def add_markdown(self, markdown: str, chapters: dict[int, str] | None = None) -> None:
        """chapters: {numero d'ordine della sezione '##': titolo del capitolo}."""
        chapters = chapters or {}
        sezione = 0
        for raw in markdown.splitlines():
            line = raw.rstrip()
            if (not line.strip() or re.fullmatch(r"[-*_]{3,}", line.strip())
                    or COMMENT_RE.match(line.strip())):
                continue

            if IMG_MARKER_RE.search(line):
                resto = []
                for pezzo in IMG_MARKER_RE.split(line):
                    if pezzo.isdigit() and f"[[IMG:{pezzo}]]" in line:
                        if not self.add_image(int(pezzo)):
                            resto.append(f"[immagine {pezzo} non disponibile]")
                    elif pezzo.strip():
                        resto.append(pezzo.strip())
                if resto:
                    self._add_runs(self._paragraph("body"), " ".join(resto))
                continue

            heading = re.match(r"^(#{1,6})\s+(.*)$", line)
            if heading:
                livello = len(heading.group(1))
                if livello == 2:
                    sezione += 1
                    if sezione in chapters:
                        self.add_chapter(chapters[sezione], first=sezione == 1)
                self._heading(heading.group(2).strip(), livello)
                continue

            bullet = re.match(r"^\s*[-*•]\s+(.*)$", line)
            if bullet:
                para = self._list_paragraph("bullet", "•\t")
                self._add_runs(para, bullet.group(1).strip())
                continue

            numbered = re.match(r"^\s*(\d+)[.)]\s+(.*)$", line)
            if numbered:
                para = self._list_paragraph("number", f"{numbered.group(1)}.\t")
                self._add_runs(para, numbered.group(2).strip())
                continue

            quote = re.match(r"^\s*>\s?(.*)$", line)
            para = self._paragraph("body")
            self._add_runs(para, quote.group(1).strip() if quote else line.strip())

    # -- indice, capitoli, numeri di pagina ------------------------------------

    def add_toc(self, titolo: str = "Indice") -> None:
        """Indice automatico: Word lo compila all'apertura, con i numeri di pagina
        reali, che noi non possiamo conoscere perché non impaginiamo."""
        self._heading(titolo, 1)
        para = self._paragraph("body")
        _add_field(
            para,
            r'TOC \o "1-2" \h \z \u',
            "Indice non ancora compilato — clic destro sull'indice → «Aggiorna campo».",
            dirty=True,
        )
        self.doc.add_paragraph().paragraph_format.page_break_before = True

    def add_chapter(self, titolo: str, first: bool = False) -> None:
        self._heading(titolo, 1)
        if not first:
            self.doc.paragraphs[-1].paragraph_format.page_break_before = True

    def add_footer(self, firma: str = FIRMA, numeri: bool = True) -> None:
        """Piè di pagina: la firma c'è sempre, il numero di pagina è opzionale e va
        all'esterno — a destra sulle dispari, a sinistra sulle pari."""
        from docx.enum.text import WD_TAB_ALIGNMENT
        from docx.shared import Emu

        self.doc.settings.odd_and_even_pages_header_footer = True

        for section in self.doc.sections:
            larghezza = Emu(section.page_width - section.left_margin - section.right_margin)
            for footer, numero_a_destra in ((section.footer, True),
                                            (section.even_page_footer, False)):
                footer.is_linked_to_previous = False
                para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                for run in list(para.runs):
                    run._r.getparent().remove(run._r)
                para.paragraph_format.tab_stops.clear_all()
                para.paragraph_format.tab_stops.add_tab_stop(larghezza, WD_TAB_ALIGNMENT.RIGHT)

                if numero_a_destra:                      # pagina dispari: esterno = destra
                    para.add_run(firma)
                    if numeri:
                        para.add_run().add_tab()
                        _add_field(para, "PAGE", "1")
                elif numeri:                             # pagina pari: esterno = sinistra
                    _add_field(para, "PAGE", "2")
                    para.add_run().add_tab()
                    para.add_run(firma)
                else:
                    para.add_run(firma)

    def save(self, path: str) -> None:
        self.doc.save(path)


# ---------------------------------------------------------------------------
# 5. CLI
# ---------------------------------------------------------------------------

def read_completed(raw_path: Path) -> tuple[int | None, str]:
    """Quanti blocchi risultano gia' elaborati nel backup Markdown, e la loro coda.

    Il conteggio e' il numero del marcatore piu' alto trovato: i marcatori portano
    l'indice reale del blocco, quindi il conto resta giusto anche se il backup
    inizia con testo non marcato (scritto da una versione precedente e ripreso
    con --start-from).

    Ritorna (None, testo) se il file c'e' ma non ha nessun marcatore: i confini fra
    i blocchi non sono ricostruibili e serve --start-from.
    """
    if not raw_path.is_file():
        return 0, ""
    content = raw_path.read_text(encoding="utf-8")
    if not content.strip():
        return 0, ""

    highest = 0
    tail_start = 0
    for line in content.splitlines(keepends=True):
        match = MARKER_RE.match(line.strip())
        if match:
            highest = max(highest, int(match.group(1)))
            tail_start = content.index(line, tail_start) + len(line)

    if highest == 0:
        return None, content
    return highest, content[tail_start:]


def build_docx(markdown: str, args, worker: "GeminiWorker | None" = None) -> None:
    chapters: dict[int, str] = {}
    if args.capitoli:
        titoli = re.findall(r"^##\s+(.+)$", markdown, re.M)
        if titoli:
            try:
                worker = worker or GeminiWorker(model=args.model, fallbacks=list(FALLBACK_MODELS))
                chapters = worker.raggruppa_capitoli(titoli)
                print(f"📚 {len(chapters)} capitoli su {len(titoli)} sezioni")
            except (Exception, SystemExit) as err:  # noqa: BLE001 - meglio senza capitoli che senza documento
                print(f"   ⚠️  Capitoli non generati ({err}): documento senza divisione in capitoli.")

    builder = DocxBuilder(args.template, sorgente.cartella_immagini(args.output))
    if args.indice:
        builder.add_toc()
    builder.add_markdown(markdown, chapters=chapters)
    builder.add_footer(args.firma, numeri=args.numeri_pagina)
    builder.save(args.output)

    origin = (f"template '{args.template}'" if builder.template_used
              else "stili di default (template non trovato)")
    missing = [k for k, v in builder.styles.items() if v is None]
    print(f"\n✅ Dispensa salvata in '{args.output}' — {origin}")
    if args.indice:
        print("   ℹ️  L'indice si compila all'apertura in Word: se resta vuoto, "
              "clic destro sull'indice → «Aggiorna campo».")
    if missing:
        print(f"   ⚠️  Stili assenti nel documento, usato il fallback: {', '.join(missing)}")


def make_template(source: str, target: str) -> None:
    """Deriva un template da un documento gia' formattato: tiene stili, margini e
    impostazioni di pagina, butta via il contenuto. Il template esistente viene
    messo da parte, mai sovrascritto e basta."""
    from docx import Document

    src = Path(source)
    if not src.is_file():
        sys.exit(f"Documento di partenza non trovato: {src.resolve()}")

    tgt = Path(target)
    if tgt.is_file():
        backup = tgt.with_name(f"{tgt.stem}_backup{tgt.suffix}")
        n = 2
        while backup.exists():
            backup = tgt.with_name(f"{tgt.stem}_backup{n}{tgt.suffix}")
            n += 1
        shutil.copy2(tgt, backup)
        print(f"📦 Template precedente messo al sicuro in '{backup}'")

    doc = Document(str(src))
    clear_body(doc)
    doc.save(str(tgt))

    probe = DocxBuilder(str(tgt))
    print(f"✅ '{tgt}' rigenerato dagli stili di '{src}'")
    for key, name in probe.styles.items():
        print(f"   {key:<14} → {name or 'assente (userò il fallback)'}")


def parse_args() -> argparse.Namespace:
    ap = argparse.ArgumentParser(description="Sbobina grezza → dispensa universitaria in Word.")
    ap.add_argument("-i", "--input", default=INPUT_FILE,
                    help=f"sbobina di partenza: .txt, .md, .docx o .pdf (default: {INPUT_FILE}). "
                         "Da .docx e .pdf vengono estratte anche le immagini")
    ap.add_argument("-o", "--output", default=OUTPUT_FILE, help=f"file Word (default: {OUTPUT_FILE})")
    ap.add_argument("-t", "--template", default=TEMPLATE_FILE, help=f"template di stile (default: {TEMPLATE_FILE})")
    ap.add_argument("-m", "--model", default=MODEL_NAME, help=f"modello Gemini (default: {MODEL_NAME})")
    ap.add_argument("--models", metavar="A,B,C",
                    help="rotazione completa dei modelli, separati da virgola: si passa al "
                         "successivo quando uno esaurisce la quota giornaliera "
                         f"(default: {MODEL_NAME},{','.join(FALLBACK_MODELS)})")
    ap.add_argument("-c", "--chunk-chars", type=int, default=CHUNK_TARGET_CHARS,
                    help=f"caratteri per blocco (default: {CHUNK_TARGET_CHARS})")
    ap.add_argument("--raw", default=RAW_MARKDOWN_FILE,
                    help=f"backup Markdown incrementale (default: {RAW_MARKDOWN_FILE})")
    ap.add_argument("--resume", action="store_true",
                    help="riprende dal primo blocco non ancora elaborato, senza rifare i precedenti")
    ap.add_argument("--start-from", type=int, metavar="N",
                    help="riprende manualmente dal blocco N (per backup senza marcatori)")
    ap.add_argument("--make-template", metavar="FILE.docx",
                    help="crea template.docx dagli stili di un documento già formattato a mano")
    ap.add_argument("--only-docx", action="store_true",
                    help="genera solo il Word dal backup Markdown, senza chiamare l'API")
    ap.add_argument("--limit", type=int, metavar="N",
                    help="elabora solo N blocchi e si ferma (per provare un modello "
                         "prima di lanciare tutto)")
    ap.add_argument("--sleep", type=float, default=SLEEP_BETWEEN_CALLS, metavar="SEC",
                    help="pausa fra le chiamate, per non saturare la quota al minuto")
    ap.add_argument("--force", action="store_true",
                    help="riparte dal blocco 1 sovrascrivendo il backup esistente")
    ap.add_argument("--indice", action=argparse.BooleanOptionalAction, default=True,
                    help="indice automatico in testa al documento (default: sì)")
    ap.add_argument("--capitoli", action=argparse.BooleanOptionalAction, default=True,
                    help="divide le sezioni in capitoli, una chiamata API in più (default: sì)")
    ap.add_argument("--numeri-pagina", action=argparse.BooleanOptionalAction, default=True,
                    help="numero di pagina in basso all'esterno (default: sì)")
    ap.add_argument("--firma", default=FIRMA, help=f"testo nel piè di pagina (default: {FIRMA!r})")
    ap.add_argument("--list-models", action="store_true",
                    help="elenca gli ID API dei modelli disponibili sulla tua chiave")
    ap.add_argument("--dry-run", action="store_true",
                    help="mostra solo la suddivisione in blocchi, senza chiamare l'API")
    return ap.parse_args()


def main() -> None:
    args = parse_args()
    raw_path = Path(args.raw)

    if args.list_models:
        list_models()
        return

    if args.make_template:
        make_template(args.make_template, args.template)
        return

    # --- solo conversione: nessuna chiamata API, nessuna quota consumata ---
    if args.only_docx:
        if not raw_path.is_file() or not raw_path.read_text(encoding="utf-8").strip():
            sys.exit(f"Nessun backup da convertire: '{raw_path}' non esiste o e' vuoto.")
        build_docx(raw_path.read_text(encoding="utf-8"), args)
        return

    source = Path(args.input)
    if not source.is_file():
        sys.exit(f"File di input non trovato: {source.resolve()}")

    cartella_img = sorgente.cartella_immagini(args.output)
    text, immagini = sorgente.estrai(source, cartella_img)
    if not text:
        sys.exit(f"Il file {source} è vuoto (o non contiene testo estraibile).")
    if immagini:
        print(f"🖼️  {len(immagini)} immagini estratte in '{cartella_img.name}/'")

    chunks = chunk_text(text, target=args.chunk_chars)
    print(f"📄 {source}: {len(text):,} caratteri → {len(chunks)} blocchi "
          f"(~{args.chunk_chars} caratteri l'uno)".replace(",", "."))

    if args.dry_run:
        for i, chunk in enumerate(chunks, 1):
            preview = chunk[:90].replace("\n", " ")
            print(f"  [{i:>3}] {len(chunk):>5} car. | {preview}…")
        return

    # --- da dove ripartire -------------------------------------------------
    done_count, done_tail = read_completed(raw_path)
    start = 0

    if args.start_from is not None:
        start = max(0, args.start_from - 1)
    elif args.resume:
        if done_count is None:
            sys.exit(
                f"'{raw_path}' esiste ma non contiene marcatori di blocco (backup di una\n"
                "versione precedente dello script). Aprilo, conta i blocchi già scritti e usa:\n"
                "  python3 dispensa.py --start-from N+1"
            )
        start = done_count
    elif done_count and not args.force:
        sys.exit(
            f"'{raw_path}' contiene già {done_count} blocchi elaborati.\n"
            "  --resume   riprende da dove eri rimasto (non spreca quota)\n"
            "  --force    ricomincia da capo sovrascrivendo il backup"
        )
    elif done_count is None and not args.force:
        sys.exit(
            f"'{raw_path}' contiene già del lavoro. Usa --start-from N per riprendere,\n"
            "oppure --force per ricominciare da capo sovrascrivendolo."
        )

    if start >= len(chunks):
        print(f"Tutti i {len(chunks)} blocchi risultano già elaborati: genero solo il Word.")
        build_docx(raw_path.read_text(encoding="utf-8"), args)
        return

    if start == 0:
        if raw_path.is_file() and raw_path.read_text(encoding="utf-8").strip():
            old = raw_path.with_name(f"{raw_path.stem}_precedente{raw_path.suffix}")
            n = 2
            while old.exists():
                old = raw_path.with_name(f"{raw_path.stem}_precedente{n}{raw_path.suffix}")
                n += 1
            shutil.move(str(raw_path), str(old))
            print(f"📦 Lavoro precedente messo da parte in '{old}'")
        raw_path.write_text("", encoding="utf-8")
    else:
        print(f"↩️  Riprendo dal blocco {start + 1}: i primi {start} sono già in '{raw_path}'.")

    # contesto di continuità: coda dell'ultimo blocco già scritto
    tail = done_tail.strip()[-CONTEXT_TAIL_CHARS:] if start > 0 else ""

    if args.models:
        rotation = [m.strip() for m in args.models.split(",") if m.strip()]
        primary, fallbacks = rotation[0], rotation[1:]
    else:
        primary, fallbacks = args.model, list(FALLBACK_MODELS)
    worker = GeminiWorker(model=primary, fallbacks=fallbacks)
    stopped_at: int | None = None

    last = min(start + args.limit, len(chunks)) if args.limit else len(chunks)

    with tqdm(total=len(chunks), initial=start, desc="Rielaborazione", unit="blocco") as bar:
        for i in range(start, last):
            try:
                elaborated = worker.elaborate(chunks[i], i + 1, len(chunks),
                                              context_tail=tail, immagini=immagini)
            except QuotaExhaustedError as err:
                stopped_at = i
                tqdm.write(f"\n🛑 {err}")
                break

            with raw_path.open("a", encoding="utf-8") as fh:
                fh.write(f"{BLOCK_MARKER.format(n=i + 1)}\n{elaborated}\n\n")
            tail = elaborated[-CONTEXT_TAIL_CHARS:]
            bar.update(1)

            if args.sleep and i + 1 < last:
                time.sleep(args.sleep)

    build_docx(raw_path.read_text(encoding="utf-8"), args, worker=worker)
    print(f"   Backup Markdown: '{raw_path}'")

    if stopped_at is None and last < len(chunks):
        print(f"\n⏸️  Fermato a {last} blocchi su {len(chunks)} come richiesto da --limit.\n"
              "   Continua con: python3 dispensa.py --resume")

    if stopped_at is not None:
        print(
            f"\n⚠️  DOCUMENTO PARZIALE: {stopped_at} blocchi su {len(chunks)}.\n"
            "   Quota giornaliera esaurita su tutti i modelli della rotazione (si azzera a\n"
            "   mezzanotte, ora del Pacifico). Riprendi senza rifare nulla con:\n"
            "     python3 dispensa.py --resume\n"
            "   Oppure aggiungi altri modelli: --models modelloA,modelloB,modelloC"
        )
        sys.exit(2)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        sys.exit("\nInterrotto. Il lavoro già fatto resta nel backup: riprendi con --resume")
